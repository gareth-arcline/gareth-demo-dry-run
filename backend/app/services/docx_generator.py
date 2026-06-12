import os
import uuid
from pathlib import Path
from docx import Document
from docx.shared import Pt
from ..models.notes import GeneratedNotes
from ..core.config import settings


def create_document(notes: GeneratedNotes, user_oid: str) -> Path:
    doc = Document()

    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Calibri'

    doc.add_heading('Meeting Notes Summary', level=0)

    for topic in notes.topics:
        doc.add_heading(topic.title, level=1)
        if topic.content:
            doc.add_paragraph(topic.content)

        for sub in topic.sub_topics:
            doc.add_heading(sub.title, level=2)
            if sub.content:
                doc.add_paragraph(sub.content)

    output_dir = Path(settings.UPLOAD_FOLDER) / user_oid
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"Notes_Summary_{uuid.uuid4().hex[:8]}.docx"
    filepath = output_dir / filename
    doc.save(str(filepath))

    return filepath
